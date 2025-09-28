import pandas as pd
from docx import Document
from docx.shared import Pt
import os
from docx2pdf import convert

# ------------------- CONFIG -------------------
TEMPLATE_FOLDER = "Templates"
OUTPUT_DOCX_FOLDER = "Certificates-DOCX"
OUTPUT_PDF_FOLDER = "Certificates-PDF"
CSV_FILE = "participants.csv"

EVENT_NAME = "Your_Event_Name"   # Global AI Bootcamp 
EVENT_DATE = "Your_Event_Date" # 28-09-2025

NAME_FONT = "Segoe Script"   # Cursive font for name
NAME_SIZE = 24  # size of the name 
# --------------------------------------------

# Ensure output folders exist
os.makedirs(OUTPUT_DOCX_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_PDF_FOLDER, exist_ok=True)

# Load participant names
data = pd.read_csv(CSV_FILE, encoding='utf-8-sig').dropna()
data.columns = data.columns.str.strip()
data['Name'] = data['Name'].str.strip()

template_path = os.path.join(TEMPLATE_FOLDER, "template.docx")

def replace_placeholders_in_run(run, participant_name):
    """
    Replace placeholders in a run while preserving all other text and layout.
    """
    run_text = run.text

    # Replace {name}
    if "{name}" in run_text:
        run.text = run_text.replace("{name}", participant_name)
        run.font.name = NAME_FONT
        run.font.size = Pt(NAME_SIZE)
        run.italic = True

    # Replace {event}
    if "{event}" in run_text:
        run.text = run.text.replace("{event}", EVENT_NAME)
        run.bold = True

    # Replace {date}
    if "{date}" in run_text:
        run.text = run.text.replace("{date}", EVENT_DATE)
        run.bold = True

def generate_certificate(template_path, participant_name):
    doc = Document(template_path)

    # Replace in paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            replace_placeholders_in_run(run, participant_name)

    # Replace in tables if any
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        replace_placeholders_in_run(run, participant_name)

    # Save DOCX
    docx_file = os.path.join(OUTPUT_DOCX_FOLDER, f"{participant_name}.docx")
    doc.save(docx_file)

    # Convert to PDF
    pdf_file = os.path.join(OUTPUT_PDF_FOLDER, f"{participant_name}.pdf")
    convert(docx_file, pdf_file)

    print(f"Saved: {docx_file} & {pdf_file}")

# Generate certificates
for _, row in data.iterrows():
    generate_certificate(template_path, row['Name'])

print("\nAll certificates generated successfully!")
